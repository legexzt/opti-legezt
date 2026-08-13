import io
import os
import re
from typing import Dict, List, Any
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.graphics.shapes import Drawing, Line
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_DIR, "static", "images", "college_logo.png")
if not os.path.exists(LOGO_PATH):
    LOGO_PATH = os.path.join(BASE_DIR, "static", "images", "college_logo.jpeg")


class ExactAcademicReportGenerator:
    """
    Generates exact 15-row per section institutional reports matching the official
    Lords Institute of Engineering and Technology templates with spacious ('khula khula')
    page-filling layouts.
    """

    @staticmethod
    def generate_pdf(report_type: str, metadata: Dict[str, Any], students: List[Dict[str, Any]], stats: Dict[str, Any] = None) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=26,
            bottomMargin=26
        )

        styles = getSampleStyleSheet()
        
        college_style = ParagraphStyle(
            'CollegeTitle',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=13,
            leading=15,
            alignment=1,
            textColor=colors.black
        )
        ugc_style = ParagraphStyle(
            'UGCTitle',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=8.5,
            leading=10.5,
            alignment=1,
            textColor=colors.HexColor("#1e293b")
        )
        aicte_style = ParagraphStyle(
            'AICTETitle',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=7.5,
            leading=9.5,
            alignment=1,
            textColor=colors.HexColor("#334155")
        )
        dept_style = ParagraphStyle(
            'DeptTitle',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=10.5,
            leading=13,
            alignment=1,
            textColor=colors.black
        )
        ay_center_style = ParagraphStyle(
            'AYCenter',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=9.5,
            leading=12,
            alignment=1,
            textColor=colors.black
        )
        subject_center_style = ParagraphStyle(
            'SubjectCenter',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=9.5,
            leading=12,
            alignment=1,
            textColor=colors.black
        )
        section_title_style = ParagraphStyle(
            'SectionTitle',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=10,
            leading=13,
            alignment=1,
            textColor=colors.black
        )
        meta_class_style = ParagraphStyle(
            'MetaClass',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=9.5,
            leading=12,
            alignment=0,
            textColor=colors.black
        )
        meta_sem_style = ParagraphStyle(
            'MetaSem',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=9.5,
            leading=12,
            alignment=2,
            textColor=colors.black
        )
        cell_style = ParagraphStyle(
            'CellText',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=9,
            leading=11,
            textColor=colors.black
        )
        cell_bold = ParagraphStyle(
            'CellBold',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=9,
            leading=11,
            textColor=colors.black
        )
        cell_center = ParagraphStyle(
            'CellCenter',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=9,
            leading=11,
            alignment=1,
            textColor=colors.black
        )
        cell_center_bold = ParagraphStyle(
            'CellCenterBold',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=9,
            leading=11,
            alignment=1,
            textColor=colors.black
        )

        elements = []

        def build_header_flowable(sub_title_line: str, show_subject: bool = True):
            header_elements = []
            
            inst_name = metadata.get("institution", "LORDS INSTITUTE OF ENGINEERING AND TECHNOLOGY")
            dept_name = metadata.get("department", "Department of Computer Science and Engineering")
            
            college_info_paragraphs = [
                Paragraph(f"<b>{inst_name}</b>", college_style),
                Paragraph("(UGC Autonomous Institution)", ugc_style),
                Paragraph("Approved by AICTE | Affiliated to Osmania University | Estd.2003 | Accredited ‘A’ grade by NAAC", aicte_style),
                Paragraph(f"<b>{dept_name}</b>", dept_style)
            ]
            
            has_logo = os.path.exists(LOGO_PATH)
            if has_logo:
                try:
                    logo_img = Image(LOGO_PATH, width=54, height=54)
                    header_tbl = Table([[logo_img, college_info_paragraphs]], colWidths=[62, 461])
                    header_tbl.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                        ('ALIGN', (0,0), (0,0), 'CENTER'),
                        ('LEFTPADDING', (0,0), (-1,-1), 0),
                        ('RIGHTPADDING', (0,0), (-1,-1), 0),
                        ('TOPPADDING', (0,0), (-1,-1), 0),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 0),
                    ]))
                    header_elements.append(header_tbl)
                except Exception:
                    for p in college_info_paragraphs:
                        header_elements.append(p)
            else:
                for p in college_info_paragraphs:
                    header_elements.append(p)

            # Solid Horizontal Divider Line across the full page width
            d = Drawing(523, 4)
            d.add(Line(0, 2, 523, 2, strokeWidth=1.25, strokeColor=colors.black))
            header_elements.append(d)
            header_elements.append(Spacer(1, 4))
            
            ay = metadata.get("academic_year", "2024-25")
            course = metadata.get("course_name", "PYTHON PROGRAMING")
            
            # Center AY
            header_elements.append(Paragraph(f"<b>AY: {ay}</b>", ay_center_style))
            
            # Center Subject (if applicable)
            if show_subject:
                header_elements.append(Paragraph(f"<b>Subject: {course}</b>", subject_center_style))

            # Center Title
            header_elements.append(Paragraph(f"<b>{sub_title_line}</b>", section_title_style))
            header_elements.append(Spacer(1, 4))

            # Class: II/A (Left) and Semester: III (Right) on same line
            meta_row = [
                Paragraph(f"<b>Class: {metadata.get('class_sec', 'II/A')}</b>", meta_class_style),
                Paragraph(f"<b>Semester: {metadata.get('semester', 'III')}</b>", meta_sem_style)
            ]
            meta_tbl = Table([meta_row], colWidths=[261.5, 261.5])
            meta_tbl.setStyle(TableStyle([
                ('LEFTPADDING', (0,0), (-1,-1), 0),
                ('RIGHTPADDING', (0,0), (-1,-1), 0),
                ('TOPPADDING', (0,0), (-1,-1), 0),
                ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ]))
            header_elements.append(meta_tbl)

            return header_elements

        def get_top_15(student_list, sort_key=None, reverse=True):
            working = list(student_list)
            if sort_key == 'cgpa':
                working.sort(key=lambda s: (s.get('cgpa') if s.get('cgpa') is not None else (s.get('sgpa') or -1)), reverse=True)
            elif sort_key == 'cie':
                working.sort(key=lambda s: (s.get('cie_marks') if s.get('cie_marks') is not None else -1), reverse=True)
            elif sort_key == 'slow_cgpa':
                fail_group = [s for s in working if s.get('backlog_count', 0) > 0 or (s.get('cgpa') is None and s.get('sgpa') is None)]
                fail_group.sort(key=lambda s: (-s.get('backlog_count', 0), s.get('roll_number', '')))
                pass_group = [s for s in working if s.get('backlog_count', 0) == 0 and (s.get('cgpa') is not None or s.get('sgpa') is not None)]
                pass_group.sort(key=lambda s: (s.get('cgpa') or s.get('sgpa') or 99))
                working = fail_group + pass_group
            elif sort_key == 'slow_cie':
                working.sort(key=lambda s: (s.get('cie_marks') if s.get('cie_marks') is not None else 0))
            return working[:15]

        # Spacious row heights: 1 header row (24pt) + 15 data rows (27pt each) = 429pt
        spacious_row_heights = [24] + [27] * 15

        is_adv = (report_type == 'advanced')
        is_slow = (report_type == 'slow')
        tier_title = "Advance Learners List" if is_adv else "Slow Learners List"

        if is_adv or is_slow:
            # ================= PAGE 1: PREVIOUS SEMESTER RESULT =================
            for h_el in build_header_flowable(f"{tier_title} – Based on the Previous Semester Result", show_subject=False):
                elements.append(h_el)

            p1_students = get_top_15(students, sort_key=('cgpa' if is_adv else 'slow_cgpa'), reverse=is_adv)
            
            table_data = [[
                Paragraph("<b>S.No</b>", cell_center_bold),
                Paragraph("<b>Roll Number</b>", cell_bold),
                Paragraph("<b>Student Name</b>", cell_bold),
                Paragraph("<b>CGPA</b>", cell_center_bold)
            ]]

            for idx, s in enumerate(p1_students):
                is_fail = (s.get("backlog_count", 0) > 0) or (s.get("cgpa") is None and s.get("sgpa") is None)
                if is_slow and is_fail:
                    cgpa_str = "Fail"
                else:
                    cgpa_val = s.get("cgpa") or s.get("sgpa")
                    cgpa_str = f"{cgpa_val:.2f}" if cgpa_val is not None else "-"
                table_data.append([
                    Paragraph(f"{idx+1}.", cell_center),
                    Paragraph(str(s.get("roll_number", "")), cell_style),
                    Paragraph(str(s.get("student_name", "")), cell_style),
                    Paragraph(cgpa_str, cell_center)
                ])

            while len(table_data) <= 15:
                idx = len(table_data)
                table_data.append([
                    Paragraph(f"{idx}.", cell_center),
                    Paragraph("", cell_style),
                    Paragraph("", cell_style),
                    Paragraph("", cell_center)
                ])

            p1_table = Table(table_data, colWidths=[45, 140, 245, 90], rowHeights=spacious_row_heights)
            p1_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.7, colors.HexColor("#0f172a")),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING', (0,0), (-1,-1), 6.5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6.5),
                ('LEFTPADDING', (0,0), (-1,-1), 5),
                ('RIGHTPADDING', (0,0), (-1,-1), 5),
            ]))
            elements.append(p1_table)
            
            elements.append(Spacer(1, 16))
            sig_table = Table([["", "Signature of the faculty"]], colWidths=[340, 180])
            sig_table.setStyle(TableStyle([
                ('ALIGN', (1,0), (1,0), 'RIGHT'),
                ('FONTNAME', (0,0), (-1,-1), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 8.5),
            ]))
            elements.append(sig_table)
            elements.append(PageBreak())

            # ================= PAGE 2: FACULTY OBSERVATION =================
            for h_el in build_header_flowable(f"{tier_title} – Based on the Faculty Observation", show_subject=True):
                elements.append(h_el)

            p2_students = get_top_15(students)
            table_data2 = [[
                Paragraph("<b>S.No</b>", cell_center_bold),
                Paragraph("<b>Roll Number</b>", cell_bold),
                Paragraph("<b>Student Name</b>", cell_bold),
                Paragraph("<b>Signature of the faculty</b>", cell_center_bold)
            ]]

            for idx, s in enumerate(p2_students):
                table_data2.append([
                    Paragraph(f"{idx+1}.", cell_center),
                    Paragraph(str(s.get("roll_number", "")), cell_style),
                    Paragraph(str(s.get("student_name", "")), cell_style),
                    Paragraph("", cell_center)
                ])

            while len(table_data2) <= 15:
                idx = len(table_data2)
                table_data2.append([
                    Paragraph(f"{idx}.", cell_center),
                    Paragraph("", cell_style),
                    Paragraph("", cell_style),
                    Paragraph("", cell_center)
                ])

            p2_table = Table(table_data2, colWidths=[45, 140, 215, 120], rowHeights=spacious_row_heights)
            p2_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.7, colors.HexColor("#0f172a")),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING', (0,0), (-1,-1), 6.5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6.5),
                ('LEFTPADDING', (0,0), (-1,-1), 5),
                ('RIGHTPADDING', (0,0), (-1,-1), 5),
            ]))
            elements.append(p2_table)
            
            elements.append(Spacer(1, 16))
            elements.append(sig_table)
            elements.append(PageBreak())

            # ================= PAGE 3: CIE 1 EVALUATION =================
            for h_el in build_header_flowable(f"{tier_title} – Based on the CIE 1 Evaluation", show_subject=True):
                elements.append(h_el)

            p3_students = get_top_15(students, sort_key=('cie' if is_adv else 'slow_cie'), reverse=is_adv)
            table_data3 = [[
                Paragraph("<b>S.No</b>", cell_center_bold),
                Paragraph("<b>Roll Number</b>", cell_bold),
                Paragraph("<b>Student Name</b>", cell_bold),
                Paragraph("<b>CIE 1 MARKS</b>", cell_center_bold)
            ]]

            for idx, s in enumerate(p3_students):
                cie = s.get("cie_marks")
                cie_str = f"{cie:g}" if cie is not None else "-"
                table_data3.append([
                    Paragraph(f"{idx+1}.", cell_center),
                    Paragraph(str(s.get("roll_number", "")), cell_style),
                    Paragraph(str(s.get("student_name", "")), cell_style),
                    Paragraph(cie_str, cell_center)
                ])

            while len(table_data3) <= 15:
                idx = len(table_data3)
                table_data3.append([
                    Paragraph(f"{idx}.", cell_center),
                    Paragraph("", cell_style),
                    Paragraph("", cell_style),
                    Paragraph("", cell_center)
                ])

            p3_table = Table(table_data3, colWidths=[45, 140, 245, 90], rowHeights=spacious_row_heights)
            p3_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.7, colors.HexColor("#0f172a")),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING', (0,0), (-1,-1), 6.5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6.5),
                ('LEFTPADDING', (0,0), (-1,-1), 5),
                ('RIGHTPADDING', (0,0), (-1,-1), 5),
            ]))
            elements.append(p3_table)
            
            elements.append(Spacer(1, 16))
            elements.append(sig_table)

        else: # Comprehensive Report
            for h_el in build_header_flowable("Comprehensive 3-Tier Student Performance Report", show_subject=True):
                elements.append(h_el)

            c_headers = ["S.No", "Roll Number", "Student Name", "CIE 1", "CGPA", "Tier", "Intervention Plan"]
            c_data = [[Paragraph(f"<b>{h}</b>", cell_bold if i > 0 else cell_center_bold) for i, h in enumerate(c_headers)]]
            
            for idx, s in enumerate(students):
                tier = s.get("tier", "Average")
                tier_color = "#16a34a" if tier == "Advanced" else ("#dc2626" if tier == "Slow" else "#0284c7")
                cie = s.get("cie_marks")
                score = s.get("cgpa") or s.get("sgpa")
                
                c_data.append([
                    Paragraph(str(idx + 1), cell_center),
                    Paragraph(str(s.get("roll_number", "")), cell_style),
                    Paragraph(str(s.get("student_name", "")), cell_style),
                    Paragraph(f"{cie:g}" if cie is not None else "-", cell_center),
                    Paragraph(f"{score:.2f}" if score is not None else "-", cell_center),
                    Paragraph(f"<font color='{tier_color}'><b>{tier}</b></font>", cell_center),
                    Paragraph(str(s.get("action_plan", "-")), cell_style)
                ])

            c_table = Table(c_data, colWidths=[30, 95, 135, 45, 50, 65, 100], repeatRows=1)
            c_table.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#0f172a")),
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#f1f5f9")),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING', (0,0), (-1,-1), 3.5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 3.5),
            ]))
            elements.append(c_table)

        doc.build(elements)
        buffer.seek(0)
        return buffer.getvalue()


    @staticmethod
    def generate_docx(report_type: str, metadata: Dict[str, Any], students: List[Dict[str, Any]]) -> bytes:
        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.4)
            s.bottom_margin = Inches(0.4)
            s.left_margin = Inches(0.5)
            s.right_margin = Inches(0.5)

        is_adv = (report_type == 'advanced')
        is_slow = (report_type == 'slow')
        tier_title = "Advance Learners List" if is_adv else "Slow Learners List"

        def add_page_header(sub_title: str, show_subject: bool = True):
            inst = metadata.get("institution", "LORDS INSTITUTE OF ENGINEERING AND TECHNOLOGY")
            dept = metadata.get("department", "Department of Computer Science and Engineering")
            
            hdr_tbl = doc.add_table(rows=1, cols=2)
            hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_tbl.columns[0].width = Inches(1.1)
            hdr_tbl.columns[1].width = Inches(6.0)

            if os.path.exists(LOGO_PATH):
                try:
                    cell_logo = hdr_tbl.rows[0].cells[0]
                    p_l = cell_logo.paragraphs[0]
                    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_l.add_run().add_picture(LOGO_PATH, width=Inches(0.85))
                except Exception:
                    pass

            cell_txt = hdr_tbl.rows[0].cells[1]
            p_inst = cell_txt.paragraphs[0]
            p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_inst.paragraph_format.space_after = Pt(1)
            r1 = p_inst.add_run(inst)
            r1.bold = True
            r1.font.size = Pt(12)

            p_sub = cell_txt.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_after = Pt(1)
            r2 = p_sub.add_run("(UGC Autonomous Institution)\nApproved by AICTE | Affiliated to Osmania University | Estd.2003 | Accredited ‘A’ grade by NAAC")
            r2.font.size = Pt(8)

            p_dept = cell_txt.add_paragraph()
            p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_dept.paragraph_format.space_after = Pt(3)
            r3 = p_dept.add_run(dept)
            r3.bold = True
            r3.font.size = Pt(10)

            ay = metadata.get("academic_year", "2024-25")
            course = metadata.get("course_name", "PYTHON PROGRAMING")
            
            # Centered AY
            p_ay = doc.add_paragraph()
            p_ay.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_ay.paragraph_format.space_before = Pt(3)
            p_ay.paragraph_format.space_after = Pt(1)
            r_ay = p_ay.add_run(f"AY: {ay}")
            r_ay.bold = True
            r_ay.font.size = Pt(9.5)

            # Centered Subject (if applicable)
            if show_subject:
                p_course = doc.add_paragraph()
                p_course.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_course.paragraph_format.space_before = Pt(1)
                p_course.paragraph_format.space_after = Pt(1)
                r_c = p_course.add_run(f"Subject: {course}")
                r_c.bold = True
                r_c.font.size = Pt(9.5)

            # Centered Section Title
            p_t = doc.add_paragraph()
            p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_t.paragraph_format.space_before = Pt(2)
            p_t.paragraph_format.space_after = Pt(3)
            r_title = p_t.add_run(sub_title)
            r_title.bold = True
            r_title.font.size = Pt(10)

            # Class (Left) and Semester (Right) table right above the main table
            tbl_meta = doc.add_table(rows=1, cols=2)
            tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl_meta.columns[0].width = Inches(3.55)
            tbl_meta.columns[1].width = Inches(3.55)
            
            p_l = tbl_meta.cell(0, 0).paragraphs[0]
            p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_l.paragraph_format.space_after = Pt(3)
            r_l = p_l.add_run(f"Class: {metadata.get('class_sec', 'II/A')}")
            r_l.bold = True
            r_l.font.size = Pt(9.5)
            
            p_r = tbl_meta.cell(0, 1).paragraphs[0]
            p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_r.paragraph_format.space_after = Pt(3)
            r_r = p_r.add_run(f"Semester: {metadata.get('semester', 'III')}")
            r_r.bold = True
            r_r.font.size = Pt(9.5)

        def set_table_borders(table):
            tblPr = table._tbl.tblPr
            borders = parse_xml(r'''
                <w:tblBorders %s>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                </w:tblBorders>
            ''' % nsdecls('w'))
            tblPr.append(borders)

        def add_faculty_signature():
            p_sig = doc.add_paragraph()
            p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_sig.paragraph_format.space_before = Pt(16)
            p_sig.paragraph_format.space_after = Pt(6)
            r = p_sig.add_run("Signature of the faculty")
            r.bold = True
            r.font.size = Pt(9.5)

        # Part 1: Previous Sem (15 rows)
        add_page_header(f"{tier_title} – Based on the Previous Semester Result", show_subject=False)
        
        def get_docx_top_15(student_list, sort_key=None, reverse=True):
            working = list(student_list)
            if sort_key == 'cgpa':
                working.sort(key=lambda s: (s.get('cgpa') if s.get('cgpa') is not None else (s.get('sgpa') or -1)), reverse=True)
            elif sort_key == 'cie':
                working.sort(key=lambda s: (s.get('cie_marks') if s.get('cie_marks') is not None else -1), reverse=True)
            elif sort_key == 'slow_cgpa':
                fail_group = [s for s in working if s.get('backlog_count', 0) > 0 or (s.get('cgpa') is None and s.get('sgpa') is None)]
                fail_group.sort(key=lambda s: (-s.get('backlog_count', 0), s.get('roll_number', '')))
                pass_group = [s for s in working if s.get('backlog_count', 0) == 0 and (s.get('cgpa') is not None or s.get('sgpa') is not None)]
                pass_group.sort(key=lambda s: (s.get('cgpa') or s.get('sgpa') or 99))
                working = fail_group + pass_group
            elif sort_key == 'slow_cie':
                working.sort(key=lambda s: (s.get('cie_marks') if s.get('cie_marks') is not None else 0))
            return working[:15]

        p1_list = get_docx_top_15(students, sort_key=('cgpa' if is_adv else 'slow_cgpa'), reverse=is_adv)
        
        tbl1 = doc.add_table(rows=16, cols=4)
        tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl1)
        for row in tbl1.rows:
            row.height = Pt(24)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for j, h in enumerate(["S.No", "Roll Number", "Student Name", "CGPA"]):
            c = tbl1.rows[0].cells[j]
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)

        for i in range(15):
            row_cells = tbl1.rows[i+1].cells
            s = p1_list[i] if i < len(p1_list) else {}
            is_fail = (s.get("backlog_count", 0) > 0) or (s.get("cgpa") is None and s.get("sgpa") is None)
            if is_slow and is_fail:
                cgpa_str = "Fail"
            else:
                cgpa_val = s.get("cgpa") or s.get("sgpa")
                cgpa_str = f"{cgpa_val:.2f}" if cgpa_val is not None else ""
            row_cells[0].paragraphs[0].add_run(f"{i+1}.")
            row_cells[1].paragraphs[0].add_run(str(s.get("roll_number", "")))
            row_cells[2].paragraphs[0].add_run(str(s.get("student_name", "")))
            row_cells[3].paragraphs[0].add_run(cgpa_str)

        add_faculty_signature()
        doc.add_page_break()

        # Part 2: Faculty Observation (15 rows)
        add_page_header(f"{tier_title} – Based on the Faculty Observation", show_subject=True)
        tbl2 = doc.add_table(rows=16, cols=3)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl2)
        for row in tbl2.rows:
            row.height = Pt(24)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for j, h in enumerate(["S.No", "Roll Number", "Student Name"]):
            c = tbl2.rows[0].cells[j]
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)

        p2_list = get_docx_top_15(students)
        for i in range(15):
            row_cells = tbl2.rows[i+1].cells
            s = p2_list[i] if i < len(p2_list) else {}
            row_cells[0].paragraphs[0].add_run(f"{i+1}.")
            row_cells[1].paragraphs[0].add_run(str(s.get("roll_number", "")))
            row_cells[2].paragraphs[0].add_run(str(s.get("student_name", "")))

        add_faculty_signature()
        doc.add_page_break()

        # Part 3: CIE 1 Evaluation (15 rows)
        add_page_header(f"{tier_title} – Based on the CIE 1 Evaluation", show_subject=True)
        p3_list = get_docx_top_15(students, sort_key=('cie' if is_adv else 'slow_cie'), reverse=is_adv)
        
        tbl3 = doc.add_table(rows=16, cols=4)
        tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl3)
        for row in tbl3.rows:
            row.height = Pt(24)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for j, h in enumerate(["S.No", "Roll Number", "Student Name", "CIE 1 MARKS"]):
            c = tbl3.rows[0].cells[j]
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)

        for i in range(15):
            row_cells = tbl3.rows[i+1].cells
            s = p3_list[i] if i < len(p3_list) else {}
            cie = s.get("cie_marks")
            cie_str = f"{cie:g}" if cie is not None else ""
            row_cells[0].paragraphs[0].add_run(f"{i+1}.")
            row_cells[1].paragraphs[0].add_run(str(s.get("roll_number", "")))
            row_cells[2].paragraphs[0].add_run(str(s.get("student_name", "")))
            row_cells[3].paragraphs[0].add_run(cie_str)

        add_faculty_signature()

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

AcademicReportGenerator = ExactAcademicReportGenerator
